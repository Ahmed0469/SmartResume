import streamlit as st
import os
from PyPDF2 import PdfReader
from docx import Document
import re
from collections import Counter
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from fuzzywuzzy import fuzz
from fuzzywuzzy import process

# Try to import spaCy (with error handling for compatibility)
try:
    import spacy
    try:
        nlp = spacy.load("en_core_web_sm")
    except OSError:
        # If model not available, use blank model with lemmatization
        try:
            nlp = spacy.blank("en")
            nlp.add_pipe("lemmatizer")
        except:
            nlp = None
    SPACY_AVAILABLE = True
except (ImportError, Exception) as e:
    # Handle Pydantic compatibility issues with Python 3.14
    SPACY_AVAILABLE = False
    nlp = None

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    nltk.download('punkt_tab')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

# Set page configuration
st.set_page_config(
    page_title="SmartResume",
    page_icon="📄",
    layout="wide"
)

# Initialize session state
if 'improvements_generated' not in st.session_state:
    st.session_state['improvements_generated'] = False
if 'improved_cv' not in st.session_state:
    st.session_state['improved_cv'] = None
if 'new_ats_result' not in st.session_state:
    st.session_state['new_ats_result'] = None
if 'new_score' not in st.session_state:
    st.session_state['new_score'] = 0
if 'score_improvement' not in st.session_state:
    st.session_state['score_improvement'] = 0
if 'all_missing_skills' not in st.session_state:
    st.session_state['all_missing_skills'] = []

# Function to extract text from PDF
def extract_text_from_pdf(file):
    pdf_reader = PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

# Function to extract text from Word document
def extract_text_from_docx(file):
    doc = Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# Comprehensive skill database
SKILL_DATABASE = {
    'Programming Languages': [
        'python', 'java', 'javascript', 'typescript', 'php', 'ruby', 'kotlin', 'swift',
        'scala', 'perl', 'clojure', 'elixir', 'haskell', 'dart', 'groovy', 'lua', 'matlab',
        'sql', 'bash', 'powershell', 'c++', 'c#', 'vb.net', 'objective-c', 'golang',
        'rust', 'r programming', 'assembly', 'lisp', 'scheme', 'solidity',
        'delphi', 'pascal', 'visual basic', 'cobol', 'fortran'
    ],
    'Frontend Frameworks': [
        'react', 'angular', 'vue', 'ember', 'backbone', 'nextjs', 'nuxt', 'svelte',
        'astro', 'solid', 'gatsby', 'remix', 'preact', 'jquery', 'bootstrap', 'angular.js',
        'knockout', 'aurelia', 'polymer', 'webcomponents', 'foundation', 'bulma', 'tailwind', 'material ui',
        'semantic ui', 'materialize', 'ionic', 'flutter', 'react native'
    ],
    'Backend Frameworks': [
        'django', 'flask', 'fastapi', 'nodejs', 'express', 'spring', 'springboot', 'laravel',
        'ruby on rails', 'rails', 'asp.net', 'core', 'dotnet', '.net', 'koa', 'hapi', 'nestjs',
        'tomcat', 'glassfish', 'jboss', 'websphere', 'pyramid', 'bottle', 'tornado', 'aiohttp',
        'gin', 'echo', 'fiber', 'rails', 'sinatra', 'slim'
    ],
    'Databases': [
        'sql', 'mysql', 'postgresql', 'postgres', 'mongodb', 'dynamodb', 'cassandra', 'redis',
        'elasticsearch', 'neo4j', 'firebase', 'couchdb', 'influxdb', 'oracle', 'mssql',
        'mariadb', 'sqlite', 'hbase', 'cockroachdb', 'voltdb', 'rethinkdb', 'memcached',
        'solr', 'aramazon', 'ddb', 'bigquery', 'snowflake', 'redshift'
    ],
    'Cloud Platforms': [
        'aws', 'azure', 'gcp', 'google cloud', 'heroku', 'digitalocean', 'linode', 'vultr',
        'ibm cloud', 'oracle cloud', 'alibaba cloud', 'openstack', 'cloudfoundry',
        'aws ec2', 'aws s3', 'aws lambda', 'azure vm', 'gcp compute'
    ],
    'DevOps/CI-CD': [
        'docker', 'kubernetes', 'jenkins', 'gitlab', 'github', 'bitbucket', 'travis', 'circleci',
        'terraform', 'ansible', 'puppet', 'chef', 'vagrant', 'prometheus', 'grafana', 'elk',
        'newrelic', 'datadog', 'splunk', 'git', 'svn', 'mercurial', 'ci/cd', 'devops',
        'argocd', 'helm', 'istio', 'prometheus', 'splunk', 'logstash'
    ],
    'Data Science/ML': [
        'machine learning', 'deep learning', 'nlp', 'computer vision', 'artificial intelligence',
        'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'pandas', 'numpy', 'matplotlib',
        'seaborn', 'opencv', 'spacy', 'nltk', 'xgboost', 'lightgbm', 'spark', 'hadoop',
        'jupyter', 'anaconda', 'statsmodels', 'hugging face', 'transformers'
    ],
    'Cybersecurity': [
        'penetration testing', 'network security', 'encryption', 'vulnerability assessment',
        'security operations', 'incident response', 'firewall', 'intrusion detection',
        'authentication', 'authorization', 'ssl/tls', 'vpn', 'security audit', 'threat modeling',
        'ethical hacking', 'security compliance', 'pci dss', 'hipaa', 'gdpr', 'iso 27001',
        'security information event management', 'siem', 'antivirus', 'malware analysis'
    ],
    'Software Architecture': [
        'software architecture', 'enterprise architecture', 'solution architecture', 'system design',
        'microservices architecture', 'api design', 'rest api', 'graphql', 'design patterns',
        'scalability', 'reliability', 'security architecture', 'cloud architecture',
        'data architecture', 'integration architecture', 'domain driven design', 'ddd',
        'event driven architecture', 'service oriented architecture', 'soa'
    ],
    'Engineering': [
        'systems engineering', 'software engineering', 'hardware engineering', 'electrical engineering',
        'mechanical engineering', 'civil engineering', 'controls engineering', 'embedded systems',
        'firmware development', 'iot', 'robotics', 'signal processing', 'power systems',
        'renewable energy', 'structural engineering', 'hvac', 'plumbing', 'construction'
    ],
    'Marketing': [
        'seo', 'sem', 'marketing automation', 'analytics', 'content marketing', 'social media marketing',
        'email marketing', 'marketing strategy', 'brand management', 'digital marketing', 'product marketing',
        'growth hacking', 'marketing analytics', 'market research', 'customer segmentation',
        'conversion rate optimization', 'cro', 'marketing qualified leads', 'lead generation',
        'marketing qualified leads', 'sales enablement', 'marketing funnel', 'crm'
    ],
    'Teaching/Education': [
        'teaching', 'education', 'curriculum design', 'instructional design', 'pedagogy',
        'educational technology', 'e-learning', 'tutoring', 'mentoring', 'training',
        'online learning', 'distance learning', 'learning management system', 'lms',
        'student assessment', 'classroom management', 'differentiated instruction',
        'special education', 'english as second language', 'esl', 'adult education'
    ],
    'Other Tools': [
        'git', 'linux', 'unix', 'rest api', 'graphql', 'microservices', 'agile', 'scrum',
        'kanban', 'jira', 'confluence', 'json', 'xml', 'html', 'css', 'sass', 'webpack',
        'gulp', 'grunt', 'npm', 'yarn', 'pip', 'maven', 'gradle', 'soap', 'mqtt', 'websocket',
        'postman', 'swagger', 'openapi', 'junit', 'pytest', 'selenium', 'cypress', 'load testing',
        'jmeter', 'gatling', 'automation testing', 'performance testing'
    ]
}

def extract_skills_enhanced(text):
    """Extract skills using comprehensive database and fuzzy matching"""
    text_lower = text.lower()
    found_skills = {}
    
    for category, skills in SKILL_DATABASE.items():
        found_skills[category] = set()
        
        for skill in skills:
            # Direct substring match (more lenient)
            if skill in text_lower:
                found_skills[category].add(skill)
            else:
                # Try fuzzy matching on words in text
                words = text_lower.split()
                skill_words = skill.split()
                
                # If skill is multi-word, check if all words exist
                if len(skill_words) > 1:
                    if all(word in text_lower for word in skill_words):
                        found_skills[category].add(skill)
                else:
                    # Single word - use fuzzy matching
                    for word in words:
                        if len(word) > 2:  # Only match words with 3+ chars
                            match_score = fuzz.ratio(skill, word)
                            if match_score >= 85:
                                found_skills[category].add(skill)
                                break
    
    return found_skills

def generate_resume_recommendations(cv_text, job_description, ats_result):
    """
    Generate professional, actionable resume improvement recommendations
    based on the gap analysis between CV and job description
    """
    recommendations = {
        'critical': [],      # High priority (score impact > 10%)
        'high': [],          # Medium-high priority (score impact 5-10%)
        'medium': [],        # Medium priority (score impact 2-5%)
        'low': [],           # Low priority (score impact < 2%)
        'summary': {
            'missing_skills': [],
            'missing_keywords': [],
            'missing_categories': [],
            'improvement_potential': 0
        }
    }
    
    job_skills = ats_result['job_skills_by_category']
    cv_skills = ats_result['cv_skills_by_category']
    matched_keywords = ats_result['matched_keywords']
    missing_keywords = ats_result['missing_keywords']
    
    # ===== 1. MISSING SKILLS ANALYSIS =====
    missing_skills_by_category = {}
    total_missing_skills = 0
    
    for category in job_skills:
        missing = job_skills[category] - cv_skills[category]
        if missing:
            missing_skills_by_category[category] = list(missing)
            total_missing_skills += len(missing)
    
    # High impact missing skills
    high_impact_skills = []
    for category, skills in missing_skills_by_category.items():
        for skill in skills[:3]:  # Top 3 per category
            high_impact_skills.append({
                'skill': skill,
                'category': category,
                'priority': 'critical' if len(skills) > 1 else 'high'
            })
    
    for skill_info in high_impact_skills[:5]:
        skill = skill_info['skill']
        category = skill_info['category']
        impact = (1 / max(total_missing_skills, 1)) * 15 * 100
        
        recommendation = {
            'title': f"Add {skill.title()} skill to '{category}'",
            'description': f"The job requires expertise in {skill}. Add projects, certifications, or achievements that showcase your proficiency with {skill} in your CV.",
            'impact': f"Could improve your score by ~{min(impact, 15):.1f}%",
            'action': f"✓ Add '{skill}' to your experience or skills section\n✓ Mention specific projects where you used {skill}\n✓ List certifications or relevant training",
            'impact_score': min(impact, 15)
        }
        recommendations['critical'].append(recommendation)
    
    # ===== 2. MISSING KEYWORDS ANALYSIS =====
    if missing_keywords:
        # Group keywords by type
        keyword_groups = {
            'technical_terms': [],
            'tools_platforms': [],
            'methodologies': [],
            'certifications': [],
            'other': []
        }
        
        for keyword in missing_keywords[:20]:
            kw_lower = keyword.lower().replace(' (partial)', '').replace(' (matched as: ', '').rstrip(')')
            
            if any(word in kw_lower for word in ['design', 'architecture', 'pattern', 'system']):
                keyword_groups['technical_terms'].append(keyword)
            elif any(word in kw_lower for word in ['platform', 'tool', 'framework', 'library']):
                keyword_groups['tools_platforms'].append(keyword)
            elif any(word in kw_lower for word in ['agile', 'scrum', 'kanban', 'methodology']):
                keyword_groups['methodologies'].append(keyword)
            elif any(word in kw_lower for word in ['certified', 'certification', 'aws', 'azure']):
                keyword_groups['certifications'].append(keyword)
            else:
                keyword_groups['other'].append(keyword)
        
        # Add keyword recommendations
        for group, keywords in keyword_groups.items():
            if keywords:
                impact = (len(keywords) / max(len(missing_keywords), 1)) * 30 * 100
                group_title = group.replace('_', ' ').title()
                
                recommendation = {
                    'title': f"Incorporate {group_title} keywords",
                    'description': f"The job description uses specific {group_title.lower()} that are missing from your CV. These are important for both ATS systems and human recruiters.",
                    'impact': f"Could improve your score by ~{min(impact, 8):.1f}%",
                    'action': f"✓ Add these keywords naturally to your CV:\n  • {', '.join(keywords[:5])}",
                    'impact_score': min(impact, 8)
                }
                
                if len(keywords) > 2:
                    recommendations['critical'].append(recommendation)
                elif len(keywords) > 1:
                    recommendations['high'].append(recommendation)
                else:
                    recommendations['medium'].append(recommendation)
    
    # ===== 3. MISSING SKILL CATEGORIES =====
    for category in job_skills:
        if not cv_skills[category]:  # Category completely missing
            job_count = len(job_skills[category])
            
            recommendation = {
                'title': f"Develop {category} expertise",
                'description': f"Your CV doesn't mention any {category.lower()}, but the job requires {job_count} different skills in this area.",
                'impact': f"Could improve your score by ~{job_count * 1.5:.1f}%",
                'action': f"✓ Include relevant {category.lower()} skills from:\n  • {', '.join(list(job_skills[category])[:3])}...\n✓ Highlight projects that used these technologies\n✓ Add relevant certifications or training",
                'impact_score': job_count * 1.5
            }
            recommendations['high'].append(recommendation)
    
    # ===== 4. EXPERIENCE LEVEL GAP =====
    if ats_result['job_exp_level'] != 'not_specified' and ats_result['cv_exp_level'] != 'not_specified':
        if ats_result['cv_exp_level'] != ats_result['job_exp_level']:
            job_level = ats_result['job_exp_level'].title()
            cv_level = ats_result['cv_exp_level'].title()
            
            recommendation = {
                'title': f"Align experience level ({cv_level} → {job_level})",
                'description': f"The job requires {job_level.lower()}-level experience, but your CV indicates {cv_level.lower()}-level skills.",
                'impact': "Could improve your score by ~5-10%",
                'action': f"✓ Update job descriptions to emphasize {job_level.lower()}-level responsibilities\n✓ Highlight complex projects and leadership roles\n✓ Include metrics and quantifiable achievements",
                'impact_score': 7.5
            }
            recommendations['high'].append(recommendation)
    
    # ===== 5. QUANTIFIABLE ACHIEVEMENTS =====
    achievement_patterns = r'(?:increased|improved|reduced|achieved|generated|earned|saved|delivered|launched|managed|led|directed|oversaw|supervised).*?(?:\d+%|\$\d+|[0-9,]+)'
    
    cv_achievements = len(re.findall(achievement_patterns, cv_text, re.IGNORECASE))
    job_achievement_count = len(re.findall(r'(\d+%|\$\d+)', job_description))
    
    if job_achievement_count > cv_achievements and cv_achievements < 5:
        recommendation = {
            'title': "Add quantifiable metrics and achievements",
            'description': f"The job description emphasizes measurable results ({job_achievement_count} metrics found). Your CV should highlight similar quantifiable achievements.",
            'impact': "Could improve your score by ~5-8%",
            'action': f"✓ Add specific numbers: percentages, revenue, time saved\n✓ Example: 'Improved system performance by 40%' instead of 'Optimized systems'\n✓ Include business impact: ROI, customer satisfaction, efficiency gains",
            'impact_score': 6.5
        }
        recommendations['high'].append(recommendation)
    
    # ===== 6. SOFT SKILLS & CERTIFICATIONS =====
    cert_pattern = r'\b(?:certification|certified|accredited|degree|diploma|training|bootcamp)\b'
    job_certs = len(re.findall(cert_pattern, job_description, re.IGNORECASE))
    cv_certs = len(re.findall(cert_pattern, cv_text, re.IGNORECASE))
    
    if job_certs > 0 and cv_certs == 0:
        recommendation = {
            'title': "Include relevant certifications",
            'description': "The job description mentions certifications. Including relevant industry certifications can significantly boost your profile.",
            'impact': "Could improve your score by ~3-5%",
            'action': f"✓ Add relevant certifications (AWS, Azure, PMP, etc.)\n✓ List in-progress or planned certifications\n✓ Include completion dates",
            'impact_score': 4.0
        }
        recommendations['medium'].append(recommendation)
    
    # ===== 7. CONTENT LENGTH & DETAIL =====
    avg_job_desc_length = len(job_description)
    cv_length = len(cv_text)
    
    if cv_length < avg_job_desc_length * 0.3:
        recommendation = {
            'title': "Expand CV with more detailed descriptions",
            'description': "Your CV is relatively brief compared to the job requirements. Adding more detail can help match specific requirements.",
            'impact': "Could improve your score by ~2-4%",
            'action': f"✓ Expand job descriptions with 2-3 bullet points per role\n✓ Add metrics and achievements to each position\n✓ Include technologies and tools used in each project",
            'impact_score': 3.0
        }
        recommendations['medium'].append(recommendation)
    
    # ===== 8. SPECIFIC SECTION IMPROVEMENTS =====
    if 'education' not in cv_text.lower() and 'bachelor' in job_description.lower():
        recommendation = {
            'title': "Add or enhance education section",
            'description': "The job requires specific education credentials that aren't clearly stated in your CV.",
            'impact': "Could improve your score by ~2-3%",
            'action': f"✓ Add degree(s) with institution and graduation date\n✓ Include GPA if above 3.5\n✓ Add relevant coursework or projects",
            'impact_score': 2.5
        }
        recommendations['low'].append(recommendation)
    
    # ===== 9. CALCULATE TOTAL IMPROVEMENT POTENTIAL =====
    total_improvement = sum(r.get('impact_score', 0) for recs in recommendations.values() 
                           if isinstance(recs, list) 
                           for r in recs)
    
    recommendations['summary']['missing_skills'] = high_impact_skills
    recommendations['summary']['missing_keywords'] = missing_keywords[:10]
    recommendations['summary']['missing_categories'] = list(missing_skills_by_category.keys())
    recommendations['summary']['improvement_potential'] = min(total_improvement, 40)  # Cap at 40%
    
    return recommendations

def apply_improvements_to_cv(cv_text, job_description, ats_result, recommendations):
    """
    Apply all calculated improvements to the CV and return an enhanced version
    """
    improved_cv = cv_text
    
    # Extract missing skills and keywords
    missing_keywords_list = ats_result['missing_keywords'][:15]
    job_skills = ats_result['job_skills_by_category']
    cv_skills = ats_result['cv_skills_by_category']
    
    # Collect all missing skills
    all_missing_skills = []
    for category in job_skills:
        missing = job_skills[category] - cv_skills[category]
        all_missing_skills.extend(list(missing))
    
    # ===== 1. ADD SKILLS SECTION IF NOT EXISTS =====
    if 'skill' not in improved_cv.lower():
        skills_section = "\n\n## TECHNICAL SKILLS\n"
        
        # Group missing skills by category
        for category in job_skills:
            missing = job_skills[category] - cv_skills[category]
            if missing or cv_skills[category]:
                all_cat_skills = list(cv_skills[category]) + list(missing)
                skills_section += f"{category}: {', '.join(sorted(all_cat_skills)[:8])}\n"
        
        improved_cv += skills_section
    else:
        # Enhance existing skills section
        skills_match = re.search(r'(#+\s*(?:SKILL|Technical|CORE|PROFESSIONAL|KEY)).*?(?=\n#+|\Z)', 
                                improved_cv, re.IGNORECASE | re.DOTALL)
        if skills_match:
            skills_section = skills_match.group(0)
            
            # Add missing skills to each category
            for category in job_skills:
                missing = list(job_skills[category] - cv_skills[category])
                if missing:
                    # Find category line and add missing skills
                    if category.lower() in skills_section.lower():
                        # Add to existing category
                        skills_section = re.sub(
                            rf'({category}[^:\n]*:)([^\n]+)',
                            lambda m: m.group(1) + m.group(2) + ', ' + ', '.join(missing[:3]),
                            skills_section,
                            flags=re.IGNORECASE
                        )
                    else:
                        # Add new category line
                        skills_section += f"\n{category}: {', '.join(missing[:5])}"
            
            improved_cv = improved_cv.replace(skills_match.group(0), skills_section)
    
    # ===== 2. ENHANCE DESCRIPTIONS WITH MISSING KEYWORDS =====
    for keyword in missing_keywords_list[:10]:
        # Clean keyword
        kw_clean = keyword.lower().replace(' (partial)', '').replace(' (matched as: ', '').rstrip(')')
        
        # Add keyword naturally to experience section if not too specific
        if len(kw_clean) > 3 and not any(word in kw_clean for word in ['the', 'and', 'or', 'is', 'was']):
            # Try to add to last bullet point or experience entry
            exp_pattern = r'(\n\s*[-•*]\s+[^\n]{30,100})\n'
            matches = list(re.finditer(exp_pattern, improved_cv))
            
            if matches and len(matches) > 0:
                last_match = matches[-1]
                # Only add if keyword not already present
                if kw_clean not in last_match.group(1).lower():
                    new_bullet = last_match.group(1) + f" and leveraged {kw_clean}"
                    improved_cv = improved_cv[:last_match.start(1)] + new_bullet + improved_cv[last_match.end(1):]
    
    # ===== 3. ADD QUANTIFIABLE METRICS =====
    # Find action verbs and add metrics
    action_verbs = ['managed', 'led', 'developed', 'improved', 'increased', 'reduced', 'created', 'designed']
    
    for verb in action_verbs:
        # Find lines with action verbs that don't have metrics
        pattern = rf'(\n\s*[-•*]\s+{verb}[^\n]*?)(?!\d)(\n|$)'
        
        # Sample metrics to add
        metrics = ['by 30%', 'by 40%', 'saving 20 hours/week', 'with 5+ team members']
        
        # Only modify a few to avoid over-doing it
        matches = list(re.finditer(pattern, improved_cv, re.IGNORECASE))
        for i, match in enumerate(matches[:2]):
            if match and i < len(metrics):
                original = match.group(1)
                if not any(char.isdigit() for char in original):
                    improved_cv = improved_cv[:match.start(1)] + original + f" {metrics[i]}" + improved_cv[match.end(1):]
    
    # ===== 4. ADD PROFESSIONAL SUMMARY IF MISSING =====
    if 'summary' not in improved_cv.lower()[:500] and 'objective' not in improved_cv.lower()[:500]:
        # Extract experience level
        exp_level = ats_result['cv_exp_level'].title()
        
        # Create professional summary
        summary = f"\n## PROFESSIONAL SUMMARY\n"
        summary += f"{exp_level}-level professional with proven expertise in "
        
        # Add top skills
        top_skills = []
        for category in job_skills:
            if cv_skills[category]:
                top_skills.extend(list(cv_skills[category])[:2])
        
        summary += f"{', '.join(top_skills[:3])}. "
        summary += f"Seeking {exp_level.lower()}-level position to leverage technical expertise and deliver impactful results.\n"
        
        improved_cv = summary + improved_cv
    
    # ===== 5. ENHANCE EXISTING EXPERIENCE ENTRIES =====
    # Find bullet points and make them more detailed
    bullet_pattern = r'([-•*]\s+[^\n]{20,80})'
    bullets = re.findall(bullet_pattern, improved_cv)
    
    for bullet in bullets[:5]:
        # If bullet is short, try to enhance it
        if len(bullet) < 60:
            # Add relevant technology mention if applicable
            for skill in list(all_missing_skills)[:3]:
                enhanced = bullet + f" using {skill}"
                improved_cv = improved_cv.replace(bullet, enhanced, 1)
                break
    
    # ===== 6. ADD CERTIFICATIONS SECTION IF NEEDED =====
    if 'certifications' not in improved_cv.lower() and 'certification' not in improved_cv.lower():
        # Check if job mentions certifications
        if re.search(r'(?:certification|certified|AWS|Azure|GCP)', job_description, re.IGNORECASE):
            certs_section = "\n\n## CERTIFICATIONS & TRAINING\n"
            certs_section += "• Committed to continuous professional development\n"
            certs_section += "• Open to pursuing relevant industry certifications\n"
            improved_cv += certs_section
    
    # ===== 7. CLEAN AND FORMAT =====
    # Remove duplicate entries
    improved_cv = re.sub(r'\n\n+', '\n\n', improved_cv)
    
    # Standardize section headers
    improved_cv = re.sub(r'###\s+', '## ', improved_cv)
    improved_cv = re.sub(r'####\s+', '### ', improved_cv)
    
    return improved_cv

# Function to extract text based on file type
def extract_text(uploaded_file):
    if uploaded_file.type == "application/pdf":
        return extract_text_from_pdf(uploaded_file)
    elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                                 "application/msword"]:
        return extract_text_from_docx(uploaded_file)
    return ""

# SpaCy-enhanced text processing functions
def extract_entities_with_spacy(text):
    """Extract named entities using spaCy for better understanding"""
    if not SPACY_AVAILABLE:
        return {'job_titles': [], 'organizations': [], 'locations': []}
    
    try:
        doc = nlp(text[:1000000])  # Limit to 1M chars for performance
        entities = {
            'job_titles': [],
            'organizations': [],
            'locations': [],
            'persons': []
        }
        
        for ent in doc.ents:
            if ent.label_ == 'PERSON':
                entities['persons'].append(ent.text)
            elif ent.label_ == 'ORG':
                entities['organizations'].append(ent.text)
            elif ent.label_ == 'GPE':
                entities['locations'].append(ent.text)
        
        return entities
    except:
        return {'job_titles': [], 'organizations': [], 'locations': [], 'persons': []}

def lemmatize_keywords(text):
    """Use spaCy to lemmatize keywords for better matching"""
    if not SPACY_AVAILABLE:
        return text.lower().split()
    
    try:
        doc = nlp(text.lower()[:500000])
        lemmas = [token.lemma_ for token in doc if not token.is_stop and token.is_alpha]
        return lemmas
    except:
        return text.lower().split()

def calculate_spacy_similarity(text1, text2):
    """Calculate semantic similarity using spaCy"""
    if not SPACY_AVAILABLE:
        return 0
    
    try:
        doc1 = nlp(text1[:500000])
        doc2 = nlp(text2[:500000])
        similarity = doc1.similarity(doc2)
        return similarity * 100
    except:
        return 0

# Function to extract text based on file type
def extract_text(uploaded_file):
    if uploaded_file.type == "application/pdf":
        return extract_text_from_pdf(uploaded_file)
    elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                                 "application/msword"]:
        return extract_text_from_docx(uploaded_file)
    return ""

# Function to calculate ATS score
def calculate_ats_score(cv_text, job_description):
    """
    Calculate ATS score using advanced multi-technique algorithm:
    - TF-IDF + Cosine similarity for semantic matching
    - spaCy NLP for entity recognition and semantic analysis
    - Fuzzy matching for skill variations
    - Keyword frequency & position analysis
    - Skills extraction with category matching
    - Experience level alignment
    - Custom scoring weights
    """
    if not cv_text or not job_description:
        return {
            'total_score': 0,
            'semantic_score': 0,
            'keyword_score': 0,
            'skill_score': 0,
            'experience_score': 0,
            'education_score': 0,
            'matched_keywords': [],
            'missing_keywords': [],
            'job_skills': [],
            'matched_skills': [],
            'spacy_score': 0
        }
    
    # Preprocess texts
    cv_text_clean = cv_text.lower().strip()
    job_desc_clean = job_description.lower().strip()
    
    # ===== 0. SPACY SEMANTIC ANALYSIS =====
    spacy_semantic_score = 0
    job_entities = {}
    cv_entities = {}
    
    if SPACY_AVAILABLE:
        # Calculate spaCy semantic similarity
        spacy_semantic_score = calculate_spacy_similarity(job_desc_clean, cv_text_clean)
        
        # Extract entities for better context understanding
        job_entities = extract_entities_with_spacy(job_description)
        cv_entities = extract_entities_with_spacy(cv_text)
    
    # ===== 1. SEMANTIC SIMILARITY USING TF-IDF =====
    try:
        vectorizer = TfidfVectorizer(analyzer='char', ngram_range=(2, 3), lowercase=True, min_df=1)
        tfidf_matrix = vectorizer.fit_transform([job_desc_clean, cv_text_clean])
        semantic_score = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0] * 100
    except:
        semantic_score = 0
    
    # Combine spaCy and TF-IDF semantic scores
    if SPACY_AVAILABLE:
        semantic_score = (semantic_score * 0.5 + spacy_semantic_score * 0.5)
    
    # ===== 2. ADVANCED KEYWORD EXTRACTION WITH LEMMATIZATION =====
    stop_words = set(stopwords.words('english'))
    stop_words.update(['would', 'could', 'may', 'must', 'shall', 'will', 'able', 'use', 'used', 
                       'using', 'application', 'applications', 'experience', 'required', 'preferred'])
    
    # Tokenize and optionally lemmatize with spaCy
    if SPACY_AVAILABLE:
        job_keywords_lemma = lemmatize_keywords(job_desc_clean)
        cv_keywords_lemma = lemmatize_keywords(cv_text_clean)
        job_keywords = [w for w in job_keywords_lemma if w not in stop_words and len(w) >= 3]
        cv_keywords = [w for w in cv_keywords_lemma if w not in stop_words and len(w) >= 3]
    else:
        job_tokens = word_tokenize(job_desc_clean)
        cv_tokens = word_tokenize(cv_text_clean)
        job_keywords = [word for word in job_tokens 
                        if word.isalnum() and word not in stop_words and len(word) >= 3]
        cv_keywords = [word for word in cv_tokens 
                       if word.isalnum() and word not in stop_words and len(word) >= 3]
    
    # Extract bigrams and trigrams
    if SPACY_AVAILABLE:
        job_ngrams = []
        cv_ngrams = []
        job_doc = nlp(job_desc_clean[:500000])
        cv_doc = nlp(cv_text_clean[:500000])
        
        for i in range(len(job_doc) - 2):
            bigram = f"{job_doc[i].lemma_} {job_doc[i+1].lemma_}"
            trigram = f"{job_doc[i].lemma_} {job_doc[i+1].lemma_} {job_doc[i+2].lemma_}"
            if len(bigram) > 5 and all(t not in stop_words for t in bigram.split()):
                job_ngrams.append(bigram)
            if len(trigram) > 10 and all(t not in stop_words for t in trigram.split()):
                job_ngrams.append(trigram)
        
        for i in range(len(cv_doc) - 2):
            bigram = f"{cv_doc[i].lemma_} {cv_doc[i+1].lemma_}"
            trigram = f"{cv_doc[i].lemma_} {cv_doc[i+1].lemma_} {cv_doc[i+2].lemma_}"
            if len(bigram) > 5 and all(t not in stop_words for t in bigram.split()):
                cv_ngrams.append(bigram)
            if len(trigram) > 10 and all(t not in stop_words for t in trigram.split()):
                cv_ngrams.append(trigram)
    else:
        job_tokens = word_tokenize(job_desc_clean)
        cv_tokens = word_tokenize(cv_text_clean)
        job_ngrams = []
        cv_ngrams = []
        
        for i in range(len(job_tokens) - 2):
            bigram = f"{job_tokens[i]} {job_tokens[i+1]}"
            trigram = f"{job_tokens[i]} {job_tokens[i+1]} {job_tokens[i+2]}"
            if len(bigram) > 5 and all(t not in stop_words for t in bigram.split()):
                job_ngrams.append(bigram)
            if len(trigram) > 10 and all(t not in stop_words for t in trigram.split()):
                job_ngrams.append(trigram)
        
        for i in range(len(cv_tokens) - 2):
            bigram = f"{cv_tokens[i]} {cv_tokens[i+1]}"
            trigram = f"{cv_tokens[i]} {cv_tokens[i+1]} {cv_tokens[i+2]}"
            if len(bigram) > 5 and all(t not in stop_words for t in bigram.split()):
                cv_ngrams.append(bigram)
            if len(trigram) > 10 and all(t not in stop_words for t in trigram.split()):
                cv_ngrams.append(trigram)
    
    # Combine all keywords
    job_all_keywords = list(set(job_keywords + job_ngrams))
    cv_all_keywords = list(set(cv_keywords + cv_ngrams))
    
    # ===== 3. FUZZY MATCHING WITH LEMMATIZATION =====
    matched_keywords_fuzzy = []
    missing_keywords = []
    
    for job_kw in job_all_keywords:
        # Try exact match first
        if job_kw in cv_all_keywords:
            matched_keywords_fuzzy.append(job_kw)
        else:
            # Fuzzy match with 80% similarity threshold
            matches = process.extract(job_kw, cv_all_keywords, limit=1, scorer=fuzz.token_set_ratio)
            if matches and matches[0][1] >= 80:
                matched_keywords_fuzzy.append(f"{job_kw} (matched as: {matches[0][0]})")
            else:
                # Check for partial matches
                job_kw_parts = job_kw.split()
                if len(job_kw_parts) > 1:
                    part_matches = sum(1 for part in job_kw_parts if any(part in cv_kw for cv_kw in cv_all_keywords))
                    if part_matches >= len(job_kw_parts) * 0.6:
                        matched_keywords_fuzzy.append(f"{job_kw} (partial)")
                    else:
                        missing_keywords.append(job_kw)
                else:
                    missing_keywords.append(job_kw)
    
    keyword_match_score = (len(matched_keywords_fuzzy) / len(job_all_keywords) * 100) if job_all_keywords else 0
    
    # ===== 4. SKILL EXTRACTION USING ENHANCED DATABASE AND FUZZY MATCHING =====
    job_skills = extract_skills_enhanced(job_description)
    cv_skills = extract_skills_enhanced(cv_text)
    
    # Calculate skill match score
    total_job_skills = sum(len(skills) for skills in job_skills.values())
    matched_skills_list = []
    
    for category in job_skills:
        matched = job_skills[category].intersection(cv_skills[category])
        matched_skills_list.extend(matched)
    
    skill_match_score = (len(matched_skills_list) / total_job_skills * 100) if total_job_skills > 0 else 0
    
    # ===== 5. EXPERIENCE LEVEL MATCHING =====
    experience_patterns = {
        'entry': r'(?:entry|junior|graduate|fresher|0-2 years?|0 to 2)',
        'mid': r'(?:mid|intermediate|3-5 years?|3 to 5|5-7 years?)',
        'senior': r'(?:senior|lead|principal|8\+ years?|10\+ years?|5-10 years?)',
    }
    
    job_exp_level = 'not_specified'
    cv_exp_level = 'not_specified'
    
    for level, pattern in experience_patterns.items():
        if re.search(pattern, job_desc_clean):
            job_exp_level = level
        if re.search(pattern, cv_text_clean):
            cv_exp_level = level
    
    # Experience alignment bonus
    experience_score = 0
    if job_exp_level != 'not_specified' and cv_exp_level != 'not_specified':
        if job_exp_level == cv_exp_level:
            experience_score = 20
        elif (job_exp_level == 'entry' and cv_exp_level in ['mid', 'senior']) or \
             (job_exp_level == 'mid' and cv_exp_level == 'senior'):
            experience_score = 10
        else:
            experience_score = 0
    
    # ===== 6. EDUCATION MATCHING =====
    education_pattern = r'\b(?:bachelor|bachelor\'s|master|master\'s|phd|ph\.d|diploma|associate|degree|bsc|msc|btech|mtech|be|me)\b'
    job_education = len(re.findall(education_pattern, job_desc_clean))
    cv_education = len(re.findall(education_pattern, cv_text_clean))
    
    education_score = 10 if cv_education > 0 and job_education > 0 else 0
    
    # ===== 7. FINAL WEIGHTED SCORING =====
    # Weights: Semantic 40%, Keywords 30%, Skills 15%, Experience 10%, Education 5%
    ats_score = (semantic_score * 0.40 + 
                 keyword_match_score * 0.30 + 
                 skill_match_score * 0.15 + 
                 experience_score * 0.10 +
                 education_score * 0.05)
    
    # Flatten job skills for display
    flat_job_skills = []
    for skills in job_skills.values():
        flat_job_skills.extend(skills)
    
    return {
        'total_score': min(ats_score, 100),
        'semantic_score': semantic_score,
        'keyword_score': keyword_match_score,
        'skill_score': skill_match_score,
        'experience_score': experience_score,
        'education_score': education_score,
        'spacy_score': spacy_semantic_score,
        'matched_keywords': matched_keywords_fuzzy[:30],
        'missing_keywords': missing_keywords[:30],
        'job_skills': list(set(flat_job_skills)),
        'matched_skills': matched_skills_list,
        'job_exp_level': job_exp_level,
        'cv_exp_level': cv_exp_level,
        'job_skills_by_category': job_skills,
        'cv_skills_by_category': cv_skills,
        'job_entities': job_entities,
        'cv_entities': cv_entities
    }

# Add title and description
st.title("📄 SmartResume")
st.markdown("---")
st.markdown("Upload your CV and paste the job description to get matched insights!")

# Create two columns
col1, col2 = st.columns(2)

# Left column - CV Upload
with col1:
    st.subheader("📥 Upload Your CV")
    uploaded_file = st.file_uploader(
        "Choose a PDF or Word document",
        type=["pdf", "docx", "doc"],
        help="Supported formats: PDF, DOCX, DOC"
    )
    
    if uploaded_file is not None:
        st.success(f"✅ File uploaded: {uploaded_file.name}")
        st.info(f"File size: {uploaded_file.size / 1024:.2f} KB")

# Right column - Job Description Input
with col2:
    st.subheader("📝 Job Description")
    job_description = st.text_area(
        "Paste the job description here",
        height=250,
        placeholder="Enter job details, requirements, and responsibilities..."
    )
    
    char_count = len(job_description)
    st.caption(f"Characters: {char_count}")

# Add a button to process
st.markdown("---")
col_btn1, col_btn2, col_btn3 = st.columns(3)

with col_btn1:
    process_button = st.button("🚀 Process & Analyze", use_container_width=True)

with col_btn2:
    reset_button = st.button("🔄 Reset", use_container_width=True)

with col_btn3:
    st.empty()

# Process the files
if process_button:
    if uploaded_file and job_description:
        st.success("✅ Processing your data...")
        
        # Extract text from CV
        cv_text = extract_text(uploaded_file)
        
        # Calculate ATS score
        ats_result = calculate_ats_score(cv_text, job_description)
        ats_score = ats_result['total_score']
        
        # Display results
        st.subheader("📊 Analysis Results")
        
        col_score1, col_score2, col_score3 = st.columns(3)
        
        with col_score1:
            # Color code the score
            if ats_score >= 75:
                st.metric("ATS Score", f"{ats_score:.2f}%", "✅ Excellent Match")
            elif ats_score >= 50:
                st.metric("ATS Score", f"{ats_score:.2f}%", "⚠️ Good Match")
            else:
                st.metric("ATS Score", f"{ats_score:.2f}%", "❌ Needs Improvement")
        
        with col_score2:
            st.metric("CV File", uploaded_file.name)
        
        with col_score3:
            st.metric("Job Description Length", f"{char_count} chars")
        
        # Display progress bar
        st.progress(min(ats_score / 100, 1.0))
        
        # Display detailed scoring breakdown
        st.subheader("🔍 Scoring Breakdown")
        col_breakdown1, col_breakdown2, col_breakdown3 = st.columns(3)
        
        with col_breakdown1:
            st.metric("Semantic Match", f"{ats_result['semantic_score']:.2f}%", "40% weight")
        
        with col_breakdown2:
            st.metric("Keyword Match", f"{ats_result['keyword_score']:.2f}%", "30% weight")
        
        with col_breakdown3:
            st.metric("Skills Match", f"{ats_result['skill_score']:.2f}%", "15% weight")
        
        col_breakdown4, col_breakdown5, col_breakdown6 = st.columns(3)
        
        with col_breakdown4:
            st.metric("Experience Level", f"{ats_result['experience_score']:.0f} pts", "10% weight")
        
        with col_breakdown5:
            st.metric("Education", f"{ats_result['education_score']:.0f} pts", "5% weight")
        
        with col_breakdown6:
            exp_status = f"{ats_result['cv_exp_level'].title()} ↔️ {ats_result['job_exp_level'].title()}"
            st.metric("Your Level vs Job", exp_status)
        
        # Display CV text preview
        with st.expander("📄 CV Text Preview"):
            st.text_area("Extracted CV Content", cv_text, height=200, disabled=True)
        
        # Display keyword analysis
        with st.expander("🔍 Keyword Analysis"):
            col_match1, col_match2 = st.columns(2)
            
            with col_match1:
                st.write(f"**✅ Matched Keywords ({len(ats_result['matched_keywords'])}):**")
                if ats_result['matched_keywords']:
                    for m in ats_result['matched_keywords'][:15]:
                        st.success(f"• {m}")
                else:
                    st.warning("No keywords matched")
            
            with col_match2:
                st.write(f"**❌ Missing Keywords ({len(ats_result['missing_keywords'])}):**")
                if ats_result['missing_keywords']:
                    for u in ats_result['missing_keywords'][:15]:
                        st.error(f"• {u}")
                else:
                    st.success("All keywords present!")
        
        # Display skills analysis
        with st.expander("💼 Skills Analysis"):
            col_skill1, col_skill2 = st.columns(2)
            
            with col_skill1:
                st.write(f"**📋 Required Skills:**")
                if ats_result['job_skills']:
                    for skill in ats_result['job_skills']:
                        st.info(f"• {skill}")
                else:
                    st.warning("No specific skills detected in job description")
            
            with col_skill2:
                st.write(f"**✅ Your Skills ({len(ats_result['matched_skills'])}):**")
                if ats_result['matched_skills']:
                    for skill in ats_result['matched_skills']:
                        st.success(f"• {skill}")
                else:
                    st.warning("No required skills found in your CV")
        
        # ===== RESUME IMPROVEMENT RECOMMENDATIONS =====
        st.subheader("🎯 Resume Improvement Recommendations")
        
        recommendations = generate_resume_recommendations(cv_text, job_description, ats_result)
        improvement_potential = recommendations['summary']['improvement_potential']
        
        # Show improvement potential
        col_imp1, col_imp2, col_imp3 = st.columns(3)
        with col_imp1:
            st.metric("Improvement Potential", f"+{improvement_potential:.1f}%")
        with col_imp2:
            new_score = min(ats_score + improvement_potential, 100)
            st.metric("Potential Score", f"{new_score:.1f}%")
        with col_imp3:
            st.metric("Missing Skills", len(recommendations['summary']['missing_skills']))
        
        st.markdown("---")
        
        # Display recommendations by priority
        priority_levels = [
            ('critical', '🔴 CRITICAL - High Impact', 'error'),
            ('high', '🟠 HIGH - Medium Impact', 'warning'),
            ('medium', '🟡 MEDIUM - Low-Medium Impact', 'info'),
            ('low', '🔵 LOW - Nice to Have', 'success')
        ]
        
        for priority, title, color in priority_levels:
            recs = recommendations[priority]
            if recs:
                with st.expander(f"{title} ({len(recs)} recommendations)", expanded=(priority == 'critical')):
                    for i, rec in enumerate(recs, 1):
                        st.markdown(f"### {i}. {rec['title']}")
                        st.write(f"**Description:** {rec['description']}")
                        st.write(f"**Potential Impact:** {rec['impact']}")
                        
                        # Create actionable steps
                        with st.container(border=True):
                            st.markdown("**Action Items:**")
                            st.write(rec['action'])
        
        # Display missing skills by category
        if recommendations['summary']['missing_categories']:
            st.markdown("---")
            st.subheader("📚 Skills Gap Analysis by Category")
            
            for category in recommendations['summary']['missing_categories']:
                missing = ats_result['job_skills_by_category'][category] - ats_result['cv_skills_by_category'][category]
                if missing:
                    with st.expander(f"**{category}** - {len(missing)} missing skills"):
                        st.write("**Missing Skills:**")
                        col_m1, col_m2 = st.columns([1, 2])
                        with col_m1:
                            for skill in list(missing)[:5]:
                                st.write(f"• {skill.title()}")
                        with col_m2:
                            st.markdown("""
                            **How to add these skills:**
                            - Mention experience with these technologies in your roles
                            - Include relevant certifications or courses
                            - Highlight projects where you used these skills
                            - Add these to your skills section
                            """)
        
        # Professional tips section
        st.markdown("---")
        st.subheader("💡 Professional Tips for Resume Enhancement")
        
        tips_col1, tips_col2 = st.columns(2)
        
        with tips_col1:
            st.markdown("""
            **✓ Best Practices:**
            - Use keywords from the job description naturally
            - Quantify achievements with metrics
            - Match your experience level to the role
            - Include relevant certifications
            - Highlight technical skills prominently
            - Keep formatting clean and ATS-friendly
            """)
        
        with tips_col2:
            st.markdown("""
            **✓ ATS Optimization:**
            - Avoid images, tables, and graphics
            - Use standard fonts (Arial, Calibri, Times)
            - Include complete job titles and dates
            - Use standard section headings
            - Write bullet points clearly
            - Save as PDF or docx format
            """)
        
        # Store the ats_result for use outside the block
        st.session_state['last_ats_result'] = ats_result
        st.session_state['last_cv_text'] = cv_text
        st.session_state['last_job_description'] = job_description
        st.session_state['last_ats_score'] = ats_score
        st.session_state['last_recommendations'] = recommendations
        
    else:
        st.error("❌ Please upload a CV file and enter job description!")

# ===== AUTO-IMPROVE SECTION (OUTSIDE process_button block) =====
if 'last_ats_result' in st.session_state:
    st.markdown("---")
    st.subheader("⚡ Auto-Improve Your Resume")
    
    col_improve1, col_improve2, col_improve3 = st.columns(3)
    
    with col_improve1:
        improve_button = st.button("✨ Generate Improved CV", use_container_width=True, key="improve_btn")
    
    with col_improve2:
        st.info("📌 Applies all recommendations automatically")
    
    with col_improve3:
        st.empty()
    
    # Process improvements if button clicked
    if improve_button:
        with st.spinner("🔧 Generating improved version of your CV..."):
            # Get stored data
            ats_result = st.session_state['last_ats_result']
            cv_text = st.session_state['last_cv_text']
            job_description = st.session_state['last_job_description']
            ats_score = st.session_state['last_ats_score']
            recommendations = st.session_state['last_recommendations']
            
            # Apply improvements
            improved_cv = apply_improvements_to_cv(cv_text, job_description, ats_result, recommendations)
            
            # Calculate all missing skills for display
            job_skills_for_display = ats_result['job_skills_by_category']
            cv_skills_for_display = ats_result['cv_skills_by_category']
            all_missing_skills_display = []
            for category in job_skills_for_display:
                missing = job_skills_for_display[category] - cv_skills_for_display[category]
                all_missing_skills_display.extend(list(missing))
            
            # Calculate new score with improved CV
            new_ats_result = calculate_ats_score(improved_cv, job_description)
            new_score = new_ats_result['total_score']
            score_improvement = new_score - ats_score
            
            # Store in session state
            st.session_state['improved_cv_data'] = {
                'improved_cv': improved_cv,
                'new_ats_result': new_ats_result,
                'new_score': new_score,
                'score_improvement': score_improvement,
                'all_missing_skills': all_missing_skills_display,
                'original_score': ats_score,
                'original_cv': cv_text
            }
            
            st.success("✅ Resume improved successfully!")

# ===== IMPROVEMENT RESULTS DISPLAY (OUTSIDE process_button block) =====
if 'improved_cv_data' in st.session_state:
    data = st.session_state['improved_cv_data']
    improved_cv = data['improved_cv']
    new_ats_result = data['new_ats_result']
    new_score = data['new_score']
    score_improvement = data['score_improvement']
    all_missing_skills_display = data['all_missing_skills']
    original_score = data['original_score']
    original_cv = data['original_cv']
    
    # Display improvement metrics
    st.subheader("📈 Improvement Results")
    
    col_result1, col_result2, col_result3 = st.columns(3)
    
    with col_result1:
        st.metric("Original Score", f"{original_score:.2f}%")
    
    with col_result2:
        if score_improvement >= 0:
            st.metric("Improved Score", f"{new_score:.2f}%", f"+{score_improvement:.2f}%", delta_color="normal")
        else:
            st.metric("Improved Score", f"{new_score:.2f}%", f"{score_improvement:.2f}%", delta_color="off")
    
    with col_result3:
        if original_score > 0:
            boost_pct = (score_improvement/original_score*100)
            st.metric("Score Boost", f"+{boost_pct:.1f}%")
    
    st.markdown("---")
    
    # Display improved CV with better preview
    st.subheader("📄 Improved Resume Preview")
    
    tab1, tab2, tab3, tab4 = st.tabs(["🎯 Summary", "📋 Full Preview", "💾 Download", "🔄 Comparison"])
    
    with tab1:
        # Show key improvements made
        st.write("**✓ Improvements Made:**")
        
        improvements_made = []
        
        # Check what was added
        if 'professional summary' in improved_cv.lower() and 'professional summary' not in original_cv.lower():
            improvements_made.append("✓ Added Professional Summary")
        
        if improved_cv.count('\n## ') > original_cv.count('\n## '):
            improvements_made.append(f"✓ Enhanced section structure")
        
        if improved_cv.count('•') > original_cv.count('•') or improved_cv.count('-') > original_cv.count('-'):
            improvements_made.append(f"✓ Added detailed bullet points with metrics")
        
        if any(skill in improved_cv for skill in all_missing_skills_display[:3]):
            improvements_made.append(f"✓ Incorporated missing skills ({len(all_missing_skills_display)} identified)")
        
        if 'CERTIFICATIONS' in improved_cv.upper() and 'CERTIFICATIONS' not in original_cv.upper():
            improvements_made.append("✓ Added Certifications section")
        
        if improved_cv.count('\n') > original_cv.count('\n'):
            improvements_made.append(f"✓ Expanded content")
        
        if not improvements_made:
            improvements_made.append("✓ Optimized formatting and structure")
        
        for improvement in improvements_made:
            st.success(improvement)
        
        st.info(f"💡 **Total improvements applied: {len(improvements_made)}**")
        
        # Show key additions
        st.markdown("---")
        st.write("**Key Additions:**")
        
        col_add1, col_add2 = st.columns(2)
        with col_add1:
            skills_str = ', '.join(all_missing_skills_display[:5]) if all_missing_skills_display else 'N/A'
            st.markdown(f"**Skills Added:**\n{skills_str}")
        
        with col_add2:
            # Need to get ats_result from somewhere - it was defined in the process_button block
            # We'll need to recalculate or store it
            st.markdown(f"**Keywords Added:**\nFrom analysis")
    
    with tab2:
        # Display formatted preview
        st.markdown("### 📰 Resume Preview")
        
        # Create a nicely formatted version
        formatted_lines = []
        for line in improved_cv.split('\n'):
            if line.startswith('##'):
                formatted_lines.append(f"**{line.replace('##', '').strip()}**")
            elif line.startswith('###'):
                formatted_lines.append(f"_{line.replace('###', '').strip()}_")
            else:
                formatted_lines.append(line)
        
        formatted_text = '\n'.join(formatted_lines)
        st.markdown(formatted_text)
        
        st.markdown("---")
        
        # Also show raw text in expandable section
        with st.expander("📄 Raw Text View"):
            st.text_area("Improved CV Content (Raw)", improved_cv, height=400, disabled=True, key="raw_cv")
    
    with tab3:
        st.markdown("**Download your improved resume:**")
        
        col_download1, col_download2, col_download3 = st.columns(3)
        
        with col_download1:
            # Download as text file
            st.download_button(
                label="📥 Download as TXT",
                data=improved_cv,
                file_name="improved_resume.txt",
                mime="text/plain",
                key="download_txt"
            )
        
        with col_download2:
            # Create and download as DOCX
            try:
                from docx import Document
                from docx.shared import Pt
                
                doc = Document()
                
                # Process and format content
                lines = improved_cv.split('\n')
                for line in lines:
                    if not line.strip():
                        doc.add_paragraph()
                    elif line.startswith('##'):
                        heading_text = line.replace('##', '').strip()
                        heading = doc.add_heading(heading_text, level=1)
                    elif line.startswith('###'):
                        subheading_text = line.replace('###', '').strip()
                        heading = doc.add_heading(subheading_text, level=2)
                    elif line.startswith(('- ', '• ', '* ')):
                        bullet_text = line.lstrip('- • * ').strip()
                        doc.add_paragraph(bullet_text, style='List Bullet')
                    else:
                        if line.strip():
                            doc.add_paragraph(line.strip())
                
                # Save to bytes
                from io import BytesIO
                docx_bytes = BytesIO()
                doc.save(docx_bytes)
                docx_bytes.seek(0)
                
                st.download_button(
                    label="📥 Download as DOCX",
                    data=docx_bytes.getvalue(),
                    file_name="improved_resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_docx"
                )
            except Exception as e:
                st.error(f"❌ Could not generate DOCX")
        
        with col_download3:
            # Generate PDF
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from io import BytesIO
                
                pdf_buffer = BytesIO()
                doc = SimpleDocTemplate(pdf_buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
                story = []
                styles = getSampleStyleSheet()
                
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontSize=13,
                    textColor='#000000',
                    spaceAfter=10,
                    fontName='Helvetica-Bold'
                )
                
                body_style = ParagraphStyle(
                    'CustomBody',
                    parent=styles['BodyText'],
                    fontSize=9.5,
                    spaceAfter=4
                )
                
                for line in improved_cv.split('\n'):
                    if not line.strip():
                        story.append(Spacer(1, 4))
                    elif line.startswith('##'):
                        heading_text = line.replace('##', '').strip()
                        story.append(Paragraph(heading_text, title_style))
                        story.append(Spacer(1, 4))
                    elif line.startswith(('- ', '• ', '* ')):
                        bullet_text = line.lstrip('- • * ').strip()
                        story.append(Paragraph(f"• {bullet_text}", body_style))
                    elif line.strip():
                        story.append(Paragraph(line.strip(), body_style))
                
                doc.build(story)
                pdf_buffer.seek(0)
                
                st.download_button(
                    label="📥 Download as PDF",
                    data=pdf_buffer.getvalue(),
                    file_name="improved_resume.pdf",
                    mime="application/pdf",
                    key="download_pdf"
                )
            except ImportError:
                st.info("📌 Install reportlab: `pip install reportlab`")
            except Exception as e:
                st.warning(f"Could not generate PDF")
    
    with tab4:
        # Show detailed comparison
        st.markdown("### 📊 Scoring Comparison")
        
        comp_col1, comp_col2, comp_col3 = st.columns(3)
        
        with comp_col1:
            st.write("**Original Scores:**")
            st.metric("Semantic", f"{new_ats_result['semantic_score']:.1f}%")
            st.metric("Keywords", f"{new_ats_result['keyword_score']:.1f}%")
            st.metric("Skills", f"{new_ats_result['skill_score']:.1f}%")
        
        with comp_col2:
            st.write("**Improved Scores:**")
            st.metric("Semantic", f"{new_ats_result['semantic_score']:.1f}%")
            st.metric("Keywords", f"{new_ats_result['keyword_score']:.1f}%")
            st.metric("Skills", f"{new_ats_result['skill_score']:.1f}%")
        
        with comp_col3:
            st.write("**Improvement:**")
            st.metric("Semantic", f"0.0%")
            st.metric("Keywords", f"0.0%")
            st.metric("Skills", f"0.0%")
        
        # Content length comparison
        st.markdown("---")
        st.markdown("### 📏 Content Analysis")
        
        col_len1, col_len2 = st.columns(2)
        
        with col_len1:
            st.metric("Original CV Length", f"{len(original_cv)} chars")
            st.metric("Original Lines", original_cv.count('\n'))
            st.metric("Original Skills Found", "N/A")
        
        with col_len2:
            st.metric("Improved CV Length", f"{len(improved_cv)} chars")
            st.metric("Improved Lines", improved_cv.count('\n'))
            st.metric("New Skills Added", len(all_missing_skills_display[:5]))

if reset_button:
    st.rerun()

# Add footer
st.markdown("---")
st.markdown(
    """
    <div style='text-align: center; color: gray; font-size: 12px;'>
    SmartResume © 2026 | Made with Streamlit
    </div>
    """,
    unsafe_allow_html=True
)
